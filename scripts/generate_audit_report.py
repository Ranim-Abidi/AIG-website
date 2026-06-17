"""Generate AiG Landing project audit report as DOCX."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = Path(__file__).resolve().parent.parent / "AiG_Landing_Project_Audit_Report.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    doc = Document()

    title = doc.add_heading("AiG Landing — Project Audit & Remediation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "AIESEC Greece Landing Page\n"
        f"Report date: {date.today().strftime('%B %d, %Y')}\n"
        "Production domain: https://newsite.aiesec.gr/"
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # =========================================================================
    # 1. PROJECT AUDIT
    # =========================================================================
    add_heading(doc, "1. Project Audit", 1)

    add_heading(doc, "1.1 Overview", 2)
    doc.add_paragraph(
        "AiG Landing is a full-stack web application for AIESEC Greece. It presents four "
        "programs (Global Volunteer, Global Talent, Global Teacher, AIESEC Member), collects "
        "signup applications, creates EXPA accounts via the AIESEC auth API, and records "
        "applications in Podio CRM for local chapter tracking."
    )

    add_heading(doc, "1.2 Architecture", 2)
    doc.add_paragraph(
        "The application follows a single-page React frontend compiled by Vite and served "
        "by a Django backend in production. The API layer is a set of JSON REST endpoints "
        "under /api/. There is no separate Node server in production."
    )
    add_bullets(
        doc,
        [
            "Presentation layer: React 19 + TypeScript (Vite 7 build → templates/dist/)",
            "Application server: Django 6 (Python) — serves SPA routes and API",
            "Database: SQLite (development default; suitable for low-traffic or replaceable for production scale)",
            "Integrations: EXPA Auth API (auth.aiesec.org/users.json), Podio REST API",
            "Mapping layer: University/faculty/LC ID mappings (lc_mapping.py, complete_mapping.py)",
            "Legacy compatibility: /api/legacy-signup/ for older pipe-separated form format",
        ],
    )

    add_heading(doc, "1.3 Repository Structure", 2)
    add_bullets(
        doc,
        [
            "Landing/ — Django project (settings.py, urls.py, WSGI)",
            "backend/ — API views, validation, EXPA/Podio services, mappings",
            "AiG landing page/templates/ — React source (src/), Vite config, public assets, dist/",
            "staticfiles/ — Archive of original static images (used to restore assets)",
            "scripts/ — Utility scripts (report generator)",
            "manage.py — Django entry point",
            ".env / .env.example / .env.production.example — Environment configuration",
        ],
    )

    add_heading(doc, "1.4 Technology Stack", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technologies"
    rows = [
        ("Frontend", "React 19, TypeScript, Vite 7, React Router 7, GSAP, jsVectorMap, Font Awesome, Lato"),
        ("Backend", "Django 6, django-cors-headers, Python requests"),
        ("Integrations", "EXPA REST signup, Podio OAuth2 password grant + item APIs"),
        ("Configuration", "Environment variables via backend/env_config.py and .env file"),
        ("Tooling", "ESLint, npm, pip / requirements.txt"),
    ]
    for layer, tech in rows:
        row = table.add_row().cells
        row[0].text = layer
        row[1].text = tech

    doc.add_paragraph()

    add_heading(doc, "1.5 API Endpoints", 2)
    add_bullets(
        doc,
        [
            "GET  /api/health/ — Health check",
            "POST /api/check-email/ — Email uniqueness (delegated to EXPA at signup)",
            "POST /api/submit-global-volunteer/ — Global Volunteer (EXPA + Podio)",
            "POST /api/submit-global-talent/ — Global Talent (EXPA + Podio)",
            "POST /api/submit-global-teacher/ — Global Teacher (EXPA + Podio)",
            "POST /api/submit-member/ — AIESEC Member application (EXPA + Podio)",
            "POST /api/legacy-signup/ — Legacy forms compatibility",
            "POST /api/test-podio/ — Dev-only Podio test (DEBUG=True only)",
        ],
    )

    add_heading(doc, "1.6 Frontend Routes", 2)
    add_bullets(
        doc,
        [
            "/ — Home (hero, stats, world map, products, testimonials)",
            "/global-volunteer — Global Volunteer program + signup form",
            "/global-talent — Global Talent program + signup form",
            "/global-teacher — Global Teacher program + signup form",
            "/member — AIESEC Member application form",
        ],
    )

    add_heading(doc, "1.7 Audit Findings (Initial State)", 2)
    add_bullets(
        doc,
        [
            "Project was functional but unfinished — placeholder mappings, minimal automated tests",
            "Secrets (Podio, EXPA, Django) were hardcoded in Python source files",
            "Form validation mismatches caused HTTP 400 before reaching EXPA",
            "EXPA errors (duplicate email) surfaced as generic HTTP 500 / 'Network error'",
            "Global Volunteer frontend posted to test Podio route instead of production endpoint",
            "Image assets overwritten by placeholders; originals recovered from staticfiles/",
            "Inconsistent form UX (browser alerts, no mobile nav on product pages)",
            "Django DEBUG=True and empty ALLOWED_HOSTS — not production-ready without configuration",
            "No environment-based configuration or .gitignore for secrets at project start",
        ],
    )

    doc.add_page_break()

    # =========================================================================
    # 2. ISSUES IDENTIFIED
    # =========================================================================
    add_heading(doc, "2. Issues Identified", 1)

    sections = [
        (
            "2.1 Form Validation (HTTP 400)",
            [
                "Password rules (8+ chars, upper, lower, number) rejected weak passwords with generic 'Validation failed'",
                "Phone validation only accepted plain 10–11 digit Greek numbers; international formats (+216, etc.) failed",
                "Date of birth accepted only YYYY-MM-DD; alternate formats (MM/DD/YYYY) could fail",
                "Global Talent / Teacher faculty values were lowercase slugs; backend expected canonical names",
                "preventReason dropdown used slug values on some pages vs full text on Global Volunteer",
            ],
        ),
        (
            "2.2 EXPA Submission (HTTP 500)",
            [
                "EXPA returns HTTP 422 for duplicate email; backend previously treated this as a network error",
                "Users saw 'Failed to create EXPA account' instead of 'email already registered'",
                "person_id vs id bug — Podio tracking could fail after successful EXPA signup",
                "Member signup hardcoded country_code +30 without phone normalization",
            ],
        ),
        (
            "2.3 Routing & Integration",
            [
                "Global Volunteer form posted to /api/test-podio/ (Podio-only test) instead of /api/submit-global-volunteer/",
                "test-podio endpoint was exposed without DEBUG guard (initially)",
                "check_email_unique always returned true — duplicates only caught at EXPA signup time",
            ],
        ),
        (
            "2.4 Frontend / UX",
            [
                "Browser alert() popups for errors and success — poor user experience",
                "Member form showed success and error banners simultaneously on retry",
                "No mobile hamburger menu on product pages (only on home)",
                "No password hint or show/hide toggle on signup forms",
                "Product card 'Learn More' buttons misaligned due to varying description lengths",
                "Excessive vertical gap between stats/map section and Our Products",
                "Global Talent missing .field-icon CSS for Available Fields cards",
            ],
        ),
        (
            "2.5 Assets & Build",
            [
                "Placeholder images overwrote real assets in public/ and dist/",
                "Frontend dist/ had to be rebuilt after source changes for Django to serve updates",
                "Django runserver failed briefly due to DjangoTemplates typo in settings.py (since fixed)",
            ],
        ),
        (
            "2.6 Security",
            [
                "Podio client secret, username, and password hardcoded in multiple service files",
                "EXPA API token hardcoded in expa_service.py",
                "SSL certificate verification disabled (verify=False) on EXPA HTTPS requests",
                "Django SECRET_KEY hardcoded in settings.py",
                "DEBUG=True and empty ALLOWED_HOSTS unsuitable for public hosting",
                "No .env / .gitignore pattern for secrets",
                "Credentials exposed in repository history — rotation required before go-live",
            ],
        ),
        (
            "2.7 Deployment / Hosting",
            [
                "No production domain configuration (ALLOWED_HOSTS, CORS, CSRF) until newsite.aiesec.gr was added",
                "Vercel staging hosts frontend only — form API calls require Django backend (see Section 4)",
                "No reverse-proxy / HTTPS headers configured for production Nginx/Caddy (now added via BEHIND_PROXY)",
            ],
        ),
    ]

    for title_text, items in sections:
        add_heading(doc, title_text, 2)
        add_bullets(doc, items)

    doc.add_page_break()

    # =========================================================================
    # 3. FIXES APPLIED
    # =========================================================================
    add_heading(doc, "3. Fixes Applied", 1)

    fix_sections = [
        (
            "3.1 Backend — Validation & Normalization",
            [
                "Added normalize_form_data(), parse_phone(), parse_dob(), normalize_faculty() in validation.py",
                "Phone validation accepts international numbers (8–15 digits; +30, +216, +357, etc.)",
                "Multiple date formats normalized to YYYY-MM-DD before age validation",
                "Faculty slug aliases mapped to canonical EXPA/Podio names",
                "All opportunity submit views call normalize_form_data() before validation",
                "Member endpoint validates password, email, and phone before EXPA submission",
            ],
        ),
        (
            "3.2 Backend — EXPA Integration",
            [
                "Refactored expa_service.py with _post_expa_user() to parse EXPA JSON error bodies",
                "Duplicate email now returns HTTP 400: 'This email is already registered...'",
                "Field-level errors (email, password, phone) returned in errors object",
                "Fixed person_id extraction in views.py and legacy_views.py for Podio tracking",
                "Phone country_code passed correctly for GV, GT, Teacher, and Member flows",
                "Added expa_error_response() helper for consistent API error responses",
                "EXPA_VERIFY_SSL=True by default (configurable via .env)",
            ],
        ),
        (
            "3.3 Backend — Routing",
            [
                "Global Volunteer frontend now uses /api/submit-global-volunteer/",
                "/api/test-podio/ gated behind DEBUG=True only",
            ],
        ),
        (
            "3.4 Security Hardening",
            [
                "Created backend/env_config.py — loads secrets from .env at project root",
                "Created backend/podio_credentials.py — shared Podio OAuth credential access",
                "Removed all hardcoded credentials from expa_service.py, podio_service.py, opportunity_podio_service.py, legacy_podio_service.py, test_podio_views.py",
                "Added .env.example, .env.production.example, and root .gitignore",
                "Django settings now read SECRET_KEY, DEBUG, ALLOWED_HOSTS, CORS from environment",
                "Production security headers enabled when DEBUG=False (HSTS, secure cookies, SSL redirect)",
                "CSRF_TRUSTED_ORIGINS and BEHIND_PROXY support for HTTPS reverse proxy (Nginx/Caddy)",
            ],
        ),
        (
            "3.5 Frontend — Forms & UX",
            [
                "Unified form feedback: inline field errors + FormAlerts banners on all signup pages",
                "Removed all browser alert() popups",
                "Shared Navbar component with mobile hamburger menu on every page",
                "PasswordField component: hint text + Show/Hide toggle",
                "Client-side validatePassword() before API submit",
                "Member form resets success state on retry; displays backend field errors",
                "Global Talent & Teacher faculty/preventReason values aligned with Global Volunteer",
            ],
        ),
        (
            "3.6 Frontend — Layout & Assets",
            [
                "Product cards use flexbox — Learn More buttons aligned at same level",
                "Reduced spacing between intro/stats/map and Our Products section",
                "Restored images from staticfiles/ to public/ and dist/",
                "Added .field-icon styles to GlobalTalent.css",
                "npm run build executed after frontend changes",
            ],
        ),
        (
            "3.7 Production Domain Configuration",
            [
                "Configured for https://newsite.aiesec.gr/ as production host",
                ".env.production.example documents required production variables",
                "Django ALLOWED_HOSTS, CORS_ALLOWED_ORIGINS, CSRF_TRUSTED_ORIGINS templates for newsite.aiesec.gr",
            ],
        ),
    ]

    for title_text, items in fix_sections:
        add_heading(doc, title_text, 2)
        add_bullets(doc, items)

    add_heading(doc, "3.8 Verification Performed", 2)
    add_bullets(
        doc,
        [
            "Global Talent API: HTTP 200 with valid new email — EXPA person_id returned",
            "Duplicate email: HTTP 400 with clear message on Global Talent and Global Volunteer",
            "Django manage.py check: no issues after security and settings updates",
            "Health endpoint: HTTP 200 at /api/health/",
            "Frontend build: npm run build completed successfully",
        ],
    )

    doc.add_page_break()

    # =========================================================================
    # 4. DEPLOYMENT
    # =========================================================================
    add_heading(doc, "4. Deployment", 1)

    add_heading(doc, "4.1 Environments", 2)
    env_table = doc.add_table(rows=1, cols=3)
    env_table.style = "Table Grid"
    h = env_table.rows[0].cells
    h[0].text = "Environment"
    h[1].text = "URL / Platform"
    h[2].text = "Purpose"
    env_rows = [
        (
            "Local development",
            "http://localhost:8000",
            "Full stack — Django serves React build + API. Recommended for correct design and working forms.",
        ),
        (
            "Vite dev server (optional)",
            "http://localhost:5173",
            "Frontend hot-reload only; API proxied to localhost:8000 via expaApi.ts",
        ),
        (
            "Staging (Vercel)",
            "User-deployed Vercel URL (frontend static hosting)",
            "Preview of UI/design. Forms require a live Django backend — Vercel alone does not run Python/Django.",
        ),
        (
            "Production (planned)",
            "https://newsite.aiesec.gr/",
            "Full stack — Django + React build + API + SSL. Primary go-live target.",
        ),
    ]
    for env, url, purpose in env_rows:
        r = env_table.add_row().cells
        r[0].text = env
        r[1].text = url
        r[2].text = purpose

    doc.add_paragraph()

    add_heading(doc, "4.2 Vercel Staging Notes", 2)
    doc.add_paragraph(
        "Vercel is used for staging/preview of the React frontend. Important limitations:"
    )
    add_bullets(
        doc,
        [
            "Vercel hosts static frontend assets only — it does not run the Django backend",
            "Signup forms call /api/submit-* endpoints; on Vercel these fail unless rewrites proxy to a Django server",
            "expaApi.ts uses relative /api when not on localhost:5173 — on Vercel this hits Vercel itself, not Django",
            "For staging form testing: use localhost:8000 or point Vercel rewrites to a deployed Django API",
            "For production: deploy full Django app to newsite.aiesec.gr (VPS, PaaS, or similar) — not Vercel alone",
        ],
    )

    add_heading(doc, "4.3 Production Deployment Checklist (newsite.aiesec.gr)", 2)
    add_numbered(
        doc,
        [
            "Build frontend: cd 'AiG landing page/templates' && npm run build",
            "Copy .env.production.example to .env on server and set all secrets (rotate credentials)",
            "Set DJANGO_DEBUG=False, DJANGO_ALLOWED_HOSTS=newsite.aiesec.gr",
            "Set CORS_ALLOWED_ORIGINS=https://newsite.aiesec.gr and CSRF_TRUSTED_ORIGINS=https://newsite.aiesec.gr",
            "Set SECURE_SSL_REDIRECT=True and BEHIND_PROXY=True (behind Nginx/Caddy)",
            "pip install -r requirements.txt",
            "python manage.py collectstatic",
            "Run Django via Gunicorn/uWSGI behind Nginx with SSL certificate",
            "Point DNS for newsite.aiesec.gr to server IP",
            "Hard-refresh browser after deploy (Ctrl+Shift+R)",
        ],
    )

    add_heading(doc, "4.4 Production .env Template", 2)
    doc.add_paragraph(
        "DJANGO_DEBUG=False\n"
        "DJANGO_SECRET_KEY=<new-random-secret>\n"
        "DJANGO_ALLOWED_HOSTS=newsite.aiesec.gr\n"
        "CORS_ALLOWED_ORIGINS=https://newsite.aiesec.gr\n"
        "CSRF_TRUSTED_ORIGINS=https://newsite.aiesec.gr\n"
        "SECURE_SSL_REDIRECT=True\n"
        "BEHIND_PROXY=True\n"
        "EXPA_VERIFY_SSL=True\n"
        "+ Podio and EXPA credentials"
    )

    doc.add_page_break()

    # =========================================================================
    # 5. RECOMMENDATIONS
    # =========================================================================
    add_heading(doc, "5. Recommendations — Future Improvements", 1)

    rec_sections = [
        (
            "5.1 Security (High Priority)",
            [
                "Rotate all Podio and EXPA credentials that were ever in source code or git history",
                "Generate a new DJANGO_SECRET_KEY for production — never reuse development key",
                "Confirm .env is never committed; audit git history for leaked secrets",
                "Enable HTTPS only on newsite.aiesec.gr; keep DEBUG=False in production",
                "Consider moving Podio to app-token auth instead of password grant long-term",
            ],
        ),
        (
            "5.2 Deployment & Infrastructure",
            [
                "Deploy Django backend to newsite.aiesec.gr (VPS, Railway, Render, DigitalOcean, etc.)",
                "Keep Vercel for frontend previews only, or add vercel.json rewrites to proxy /api to Django",
                "Use PostgreSQL instead of SQLite if traffic or concurrent writes increase",
                "Set up CI/CD: build frontend + run tests on push, deploy to staging/production",
                "Configure automated SSL renewal (Let's Encrypt via Certbot or Caddy)",
                "Add uptime monitoring for /api/health/ and form submission endpoints",
            ],
        ),
        (
            "5.3 UX & Product",
            [
                "Searchable university/campus dropdown (100+ options) — major usability win",
                "Post-submit success screen with next steps and link to EXPA sign-in",
                "FAQ section per product (cost, duration, eligibility)",
                "Optional 2-step Member form with progress indicator",
            ],
        ),
        (
            "5.4 Engineering Quality",
            [
                "Add automated API tests for validation, EXPA error mapping, and happy-path signup (mocked)",
                "Add integration tests for university/faculty mapping edge cases",
                "Sanitize server logs — avoid logging full payloads with PII",
                "Document handoff README with env setup, build, and deploy steps",
            ],
        ),
        (
            "5.5 EXPA / Podio Integration",
            [
                "Verify faculty and LC placeholder IDs against live EXPA admin data",
                "Complete university mapping for any campuses missing from complete_mapping.py",
                "Handle EXPA maintenance windows with user-friendly retry messaging",
                "Monitor Podio submission failures separately from EXPA (log alerts)",
            ],
        ),
    ]

    for title_text, items in rec_sections:
        add_heading(doc, title_text, 2)
        add_bullets(doc, items)

    doc.add_paragraph()
    p = doc.add_paragraph("— End of Report —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
